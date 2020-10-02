# Say you have a text file of guest names. This guests.txt file has one name per line, as follows:
#
# Prof. Plum
# Miss Scarlet
# Col. Mustard
# Al Sweigart
# RoboCop
#
# Write a program that would generate a Word document with custom invitations that look like Figure 15-11.
# Since Python-Docx can use only those styles that already exist in the Word document,
# you will have to first add these styles to a blank Word file and then open that file with Python-Docx.
# There should be one invitation per page in the resulting Word document, so call add_break() to add a page break after the last paragraph of each invitation.
# This way, you will need to open only one Word document to print all of the invitations at once.


from docx import Document
#from docx.document import Document


# get list of guests
with open("guests.txt", "r") as guests_file:
    guests = [str(g).rstrip("\n") for g in guests_file.readlines()]

# prepare document
mydoc = Document("invitation.docx")
for guest in guests:
    mydoc.add_paragraph("It would be a pleasure to have company of", style="invitation_style")
    mydoc.add_paragraph(f"{guest}", style="invitation_guest")
    mydoc.add_paragraph("at 11010 Memory Lane at the evening of", style="invitation_style")
    mydoc.add_paragraph(f"April 1st", style="invitation_date")
    mydoc.add_paragraph("at 7 o'clock", style="invitation_style")
    mydoc.add_page_break()

mydoc.save(f"invitation_{len(guests)}_guests.docx")
print("Done")